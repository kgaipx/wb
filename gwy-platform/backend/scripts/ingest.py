"""题库导入管线（第三方参考题库 → Question 表，schema 与 load_data.py 对齐）。

支持格式：Word(.docx) 与 PDF(.pdf)。两种格式共用同一套解析逻辑。

用法（在 backend/ 目录下运行）：
    python scripts/ingest.py                 # 干跑预览：解析 incoming/*.docx/*.pdf，输出 _preview.json 并打印统计
    python scripts/ingest.py --import        # 确认无误后写入数据库（按 stem 去重，幂等）
    python scripts/ingest.py --dir path      # 指定其他目录
    python scripts/ingest.py --subject 行测 --category 言语理解与表达   # 覆盖科目/细分

解析针对常见公考刷题资料（如齐麟刷题组）布局：
- 题号开头（允许与正文粘连，如 "112.2021年…"）；选项 A./B./C./D. 常**同行内联**（"A.5 B.6 C.7 D.8"）。
- 答案通常以**独立答案区**给出（"【参考答案】CBDAB DCBAB"），按文档出现顺序逐题映射。
- 资料分析题的数据材料（表格/段落）会作为题干前缀附加到同组各题。
- 无选项且含申论关键词的段落按申论处理（qtype=essay）。
- 导入默认 is_verified=False、source=文件名、copyright_owner="导入-待核实"（第三方资料，上线前需版权校验+双签）。
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

# 允许以 `python scripts/ingest.py` 方式运行时导入 backend/app 包
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ---------- 文本抽取：Word ----------
def extract_docx_items(path: str) -> list[str]:
    try:
        import docx
    except ImportError:
        print("缺少依赖 python-docx，请先安装：pip install python-docx")
        sys.exit(1)
    d = docx.Document(path)
    lines: list[str] = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)
    return lines


# ---------- 文本抽取：PDF ----------
def extract_pdf_items(path: str) -> list[str]:
    try:
        import pdfplumber
    except ImportError:
        print("缺少依赖 pdfplumber，请先安装：pip install pdfplumber")
        sys.exit(1)
    lines: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            for ln in txt.splitlines():
                t = ln.strip()
                if t:
                    lines.append(t)
            for tbl in page.extract_tables() or []:
                for row in tbl:
                    for cell in row:
                        if cell:
                            t = str(cell).strip()
                            if t:
                                lines.append(t)
    return lines


# ---------- 正则 ----------
# 题号开头：允许与正文粘连（不强制尾部空白），但避免小数（"3.5"）——要求题号后非数字或含空白/中文
RE_Q_START = re.compile(r"^\s*(?:\(\d{1,3}\)|\d{1,3}[.．、)])")
RE_OPTION_MARKER = re.compile(r"([A-Ha-h])[.．、]\s*")
RE_KEY = re.compile(r"参\s*考\s*答\s*案")
RE_ESSAY = re.compile(r"(给定资料|作答要求|请围绕|写一篇|不少于|左右。|题目：|结合材料)")
RE_CN_NUM = re.compile(r"^[一二三四五六七八九十]+[、.．]")

SUBJECT_HINTS = {
    "行测": ["行测", "行政职业能力", "言语", "数量", "判断", "资料分析", "常识", "数资", "资料"],
    "申论": ["申论", "材料作文", "大作文"],
}


def classify_subject(text_sample: str, default: str) -> str:
    for sub, kws in SUBJECT_HINTS.items():
        if any(k in text_sample for k in kws):
            return sub
    return default


# ---------- 选项抽取（支持同行内联 A.5 B.6 C.7 D.8 / 多行） ----------
def extract_options(text: str):
    """从整段文本中抽取选项。返回 (options:list[(label,content)], stem:str)。
    选项起始以首个 'A' 起、且后续为 B,C,D 连续序列的标记为准。"""
    markers = list(RE_OPTION_MARKER.finditer(text))
    if not markers:
        return None, text.strip()
    # 找以 A 开头、且连续 B,C,D 的起始
    start_i = None
    for i, m in enumerate(markers):
        if m.group(1).upper() != "A":
            continue
        seq = [mk.group(1).upper() for mk in markers[i:i + 4]]
        if seq[: min(4, len(seq))] == ["A", "B", "C", "D"][: min(4, len(seq))]:
            start_i = i
            break
    if start_i is None:
        return None, text.strip()
    opt_markers = markers[start_i:]
    opts = []
    for j, m in enumerate(opt_markers):
        s = m.end()
        e = opt_markers[j + 1].start() if j + 1 < len(opt_markers) else len(text)
        opts.append((m.group(1).upper(), text[s:e].strip()))
    stem = text[: opt_markers[0].start()].strip()
    return opts, stem


def extract_key_letters(line: str) -> list[str]:
    """从 '【参考答案】CBDAB DCBAB' 抽取答案字母序列（逐题顺序）。"""
    # 取 '答案' 之后的部分
    m = RE_KEY.search(line)
    tail = line[m.end():] if m else line
    out: list[str] = []
    for tok in re.findall(r"[A-Ha-h]+", tail):
        if len(tok) == 1:
            out.append(tok.upper())
        else:
            # 形如 CBDAB（每字符一题）或 AB CD（每 token 一题，多选）
            # 数资刷题多为单选，逐字符展开；若整串>5且含空格分组则按 token
            out.extend(tok.upper())
    return out


def is_q_start(ln: str) -> bool:
    # 行首 N. / N、/ N) 即视为题号起点（允许与正文粘连如 "112.2021年…"）。
    # 误判（如行首 "3.5倍"）通常无 ABCD 选项，会在 flush 时被自然丢弃。
    return bool(RE_Q_START.match(ln))


NOISE_FRAGMENTS = ["目录", "微信公众号", "新浪微博", "练习说明", "内部交流",
                  "参考答案", "答案见", "视频讲解", "公考齐麟", "齐麟公考"]


def is_noise(stem: str) -> bool:
    if any(b in stem for b in NOISE_FRAGMENTS):
        return True
    if stem.count(".") > 8:
        return True
    return False


def parse_lines(lines: list[str], default_subject: str = "行测") -> tuple[list[dict], list[str]]:
    """返回 (questions, warnings)。答案按文档顺序从答案区映射。"""
    questions: list[dict] = []
    answer_pool: list[str] = []
    material: list[str] = []          # 当前数据材料（题组前的表格/段落）
    buf: list[str] = []
    fname = "doc/pdf"

    def flush_q(block: list[str]):
        if not block:
            return
        head = block[0]
        joined = "\n".join(block)
        opts, stem = extract_options(joined)
        if not opts:
            if RE_ESSAY.search(joined):
                requirement = joined.split("作答要求", 1)[1] if "作答要求" in joined else ""
                stem = re.sub(r"^\s*(?:\(\d{1,3}\)|\d{1,3}[.．、)])\s*", "", head, count=1).strip()
                questions.append({
                    "subject": "申论",
                    "category": "申论写作",
                    "qtype": "essay",
                    "stem": (stem + "\n" + joined) if stem else joined,
                    "difficulty": 3,
                    "knowledge_point": "申论",
                    "answer": requirement.strip() or None,
                    "explanation": None,
                    "options": [],
                    "source": fname,
                    "copyright_owner": "导入-待核实",
                    "is_verified": False,
                    "_fname": fname,
                })
            return
        # 仅保留有效单选题（2~4 个选项），其余（目录/页脚噪声、被误吞的多题块）丢弃，
        # 既不入库也不占用答案映射槽位，保证答案顺序与真实题目对齐。
        if len(opts) not in (2, 3, 4) or is_noise(stem):
            return
        # 去掉题干开头题号
        stem = re.sub(r"^\s*(?:\(\d{1,3}\)|\d{1,3}[.．、)])\s*", "", stem, count=1).strip()
        if material:
            stem = "\n".join(material) + "\n" + stem
        questions.append({
            "subject": classify_subject(joined, default_subject),
            "category": "待分类",
            "qtype": "single",
            "stem": stem,
            "difficulty": 3,
            "knowledge_point": "待标注",
            "answer": None,          # 稍后按序映射
            "explanation": None,
            "options": [[lbl, content, False] for lbl, content in opts],
            "source": fname,
            "copyright_owner": "导入-待核实",
            "is_verified": False,
            "_fname": fname,
        })

    for ln in lines:
        if RE_KEY.search(ln):
            ans = extract_key_letters(ln)
            if ans:
                answer_pool.extend(ans)
                material = []   # 答案区后材料重置
                buf = []
                continue
        if is_q_start(ln):
            flush_q(buf)
            buf = [ln]
            material = []       # 新题组开始，旧材料清空（如题干前紧跟新段落会重新累积）
            continue
        # 判定是否为材料行（非选项、非题干号）
        if RE_OPTION_MARKER.search(ln) and not is_q_start(ln):
            # 选项行属于当前题，进入 buf
            buf.append(ln)
        elif is_q_start(ln):
            pass
        else:
            # 可能是材料/说明；若 buf 为空则视为材料累积，否则可能是题干续行
            if not buf:
                material.append(ln)
            else:
                buf.append(ln)
    flush_q(buf)

    # 按文档顺序映射答案
    ai = 0
    unmapped = 0
    for q in questions:
        if q["qtype"] == "essay":
            continue
        if ai < len(answer_pool):
            letter = answer_pool[ai]
            # 仅当答案字母确实是本题某个选项时才赋值；否则留空（待核对），避免给出错误答案
            if any(lbl == letter for lbl, _c, _ in q["options"]):
                q["answer"] = letter
                q["options"] = [[lbl, c, lbl == letter] for lbl, c, _ in q["options"]]
            ai += 1
        else:
            unmapped += 1
    warnings = []
    if unmapped:
        warnings.append(f"{unmapped} 题未匹配到答案（答案区字母数={len(answer_pool)}，题数={len(questions)}）")
    return questions, warnings


def parse_doc(path: str, default_subject: str = "行测") -> tuple[list[dict], list[str]]:
    ext = os.path.splitext(path)[1].lower()
    lines = extract_pdf_items(path) if ext == ".pdf" else extract_docx_items(path)
    items, warns = parse_lines(lines, default_subject)
    for it in items:
        it["_fname"] = os.path.basename(path)
        it["source"] = os.path.basename(path)
    return items, warns


# ---------- 去重 + 持久化 ----------
def dedupe(items: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for it in items:
        key = it["stem"]
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


def import_to_db(items: list[dict]) -> int:
    from sqlalchemy.orm import Session
    from app.db.session import SessionLocal
    from app.models import Question, QuestionOption

    s: Session = SessionLocal()
    existing = {q.stem for q in s.query(Question.stem).all()}
    added = 0
    try:
        for it in items:
            if it["stem"] in existing:
                continue
            q = Question(
                subject=it["subject"],
                category=it.get("category", "待分类"),
                qtype=it.get("qtype", "single"),
                stem=it["stem"],
                difficulty=it.get("difficulty", 3),
                knowledge_point=it.get("knowledge_point", "待标注"),
                answer=it.get("answer"),
                explanation=it.get("explanation"),
                source=it.get("source", "导入"),
                copyright_owner=it.get("copyright_owner", "导入-待核实"),
                is_verified=bool(it.get("is_verified", False)),
            )
            q.options = [
                QuestionOption(label=lbl, content=ct, is_correct=bool(ic))
                for (lbl, ct, ic) in it.get("options", [])
            ]
            s.add(q)
            existing.add(it["stem"])
            added += 1
        s.commit()
    finally:
        s.close()
    return added


def main():
    ap = argparse.ArgumentParser(description="题库导入（Word/PDF）")
    ap.add_argument("--dir", default=os.path.join(os.path.dirname(__file__), "..", "data", "incoming"))
    ap.add_argument("--import", dest="do_import", action="store_true", help="写入数据库（默认仅预览）")
    ap.add_argument("--subject", default=None, help="覆盖 subject")
    ap.add_argument("--category", default=None, help="覆盖 category")
    ap.add_argument("--include", default="docx,pdf", help="逗号分隔的扩展名（默认 docx,pdf）")
    args = ap.parse_args()

    d = Path(args.dir)
    exts = [e.strip().lower().lstrip(".") for e in args.include.split(",") if e.strip()]
    docs: list[Path] = []
    for e in exts:
        docs.extend(d.glob(f"*.{e}"))
    if not docs:
        print(f"在 {d} 未找到 {','.join('.' + e for e in exts)} 文件。请把百度网盘下载的题库放进来。")
        return

    all_items: list[dict] = []
    for doc in docs:
        items, warns = parse_doc(str(doc), default_subject=args.subject or "行测")
        if args.subject:
            for it in items:
                it["subject"] = args.subject
        if args.category:
            for it in items:
                it["category"] = args.category
        print(f"[解析] {doc.name}: {len(items)} 题" + (f"  ⚠ {warns}" if warns else ""))
        all_items.extend(items)

    all_items = dedupe(all_items)
    preview = [{k: v for k, v in it.items() if not k.startswith("_")} for it in all_items]
    out_path = d / "_preview.json"
    out_path.write_text(json.dumps(preview, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[预览] 共 {len(preview)} 题（去重后），写入 {out_path}")
    for s in preview[:3]:
        print("  -", s["subject"], "/", s["qtype"], "|", s["stem"][:36].replace("\n", " "),
              "| 选项", len(s["options"]), "| 答案", s.get("answer"))

    if args.do_import:
        n = import_to_db(all_items)
        print(f"[导入] 新增 {n} 题到数据库（按 stem 去重，幂等）。")
    else:
        print("[提示] 未写入数据库。确认字段映射无误后加 --import 执行。")


if __name__ == "__main__":
    main()
